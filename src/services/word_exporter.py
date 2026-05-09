import io
import os
import tempfile
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from sqlalchemy.orm import Session, joinedload
from src.models.entities import PlanMacro, Policy, StrategicItem, Activity, Task, Evidence
import plotly.express as px
import pandas as pd

class WordExporterService:
    """
    Service responsible for generating Word (.docx) reports.
    """
    @staticmethod
    def _add_heading(doc, text, level, color_hex=None):
        heading = doc.add_heading(text, level=level)
        if color_hex:
            # Add custom color if provided
            for run in heading.runs:
                # Convert hex to RGB tuple
                h = color_hex.lstrip('#')
                run.font.color.rgb = RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
        return heading

    @staticmethod
    def generate_report(db: Session, start_date: datetime, end_date: datetime, filter_label: str) -> io.BytesIO:
        """
        Generates a comprehensive Word report for a specific period.
        """
        doc = Document()
        
        # --- TITLE & HEADER ---
        title = doc.add_heading('Sistema Integral de Seguimiento Estratégico de Talento Humano', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph(f'Reporte de Gestión Estratégica: {filter_label}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].bold = True
        
        gen_date = doc.add_paragraph(f'Fecha de Generación: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        gen_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()

        # --- DATA GATHERING ---
        # Fetch the Macro Plan (assuming the current active one)
        # In a multi-year setup, we might need to filter by year.
        # We will get the most recent Macro Plan that intersects with the year of the filter.
        target_year = start_date.year
        macro = db.query(PlanMacro).filter(PlanMacro.year == target_year).options(
            joinedload(PlanMacro.policies).joinedload(Policy.strategic_items).joinedload(StrategicItem.activities).joinedload(Activity.tasks)
        ).first()

        if not macro:
            # Fallback to the latest if no specific year matches
            macro = db.query(PlanMacro).order_by(PlanMacro.year.desc()).first()
            if not macro:
                doc.add_paragraph("No hay datos estratégicos registrados en el sistema.")
                bio = io.BytesIO()
                doc.save(bio)
                bio.seek(0)
                return bio

        # --- CHART GENERATION (PLOTLY TO IMAGE) ---
        doc.add_heading('1. Resumen Consolidado de Políticas', level=1)
        
        pol_data = [{"Política": p.name, "Avance (%)": p.progress} for p in macro.policies]
        if pol_data:
            df_pol = pd.DataFrame(pol_data).sort_values("Avance (%)", ascending=True)
            fig = px.bar(
                df_pol, x='Avance (%)', y='Política', orientation='h',
                title="Cumplimiento por Política Institucional",
                text_auto='.1f',
                color='Avance (%)',
                color_continuous_scale=[[0, '#ef4444'], [0.5, '#f59e0b'], [1, '#10b981']],
                template="plotly_white"
            )
            fig.update_layout(showlegend=False, xaxis=dict(range=[0, 100]))
            
            # Save to temporary file, close it, insert it, and then delete it
            # This prevents Windows [WinError 32] file lock errors.
            tmp_path = ""
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmpfile:
                tmp_path = tmpfile.name
                fig.write_image(tmp_path, engine="kaleido", width=800, height=400)
            
            # Now the file is closed, docx can read it and we can delete it safely.
            doc.add_picture(tmp_path, width=Inches(6.0))
            os.unlink(tmp_path)
        else:
            doc.add_paragraph("No hay políticas registradas para este periodo.")

        doc.add_page_break()

        # --- HIERARCHICAL DATA AND TASKS ---
        doc.add_heading('2. Detalle de Ejecución Estratégica', level=1)
        doc.add_paragraph(f"Se listan a continuación las tareas programadas o ejecutadas en el periodo: {filter_label}.")

        for pol in macro.policies:
            WordExporterService._add_heading(doc, f"Política: {pol.name} (Avance: {pol.progress:.1f}%)", level=2, color_hex="#00594e")
            
            for item in pol.strategic_items:
                WordExporterService._add_heading(doc, f"{item.type}: {item.name} (Avance: {item.progress:.1f}%)", level=3, color_hex="#3b82f6")
                
                for act in item.activities:
                    # Filter tasks for this activity that intersect with the date range
                    filtered_tasks = []
                    for t in act.tasks:
                        # Logic: task overlaps with period if it starts before period ends AND ends after period starts
                        # Or if it has no dates, we might skip it or include it if 'Cumplida' in this period
                        t_start = t.start_date
                        t_end = t.end_date
                        
                        if t_start and t_end:
                            # Date type conversion to datetime for comparison
                            t_start_dt = datetime.combine(t_start, datetime.min.time())
                            t_end_dt = datetime.combine(t_end, datetime.max.time())
                            
                            if t_start_dt <= end_date and t_end_dt >= start_date:
                                filtered_tasks.append(t)
                        elif t.fulfillment_date:
                            # If no scheduled dates but fulfilled in period
                            if start_date <= t.fulfillment_date <= end_date:
                                filtered_tasks.append(t)
                    
                    if filtered_tasks:
                        WordExporterService._add_heading(doc, f"Actividad: {act.name} (Avance: {act.progress:.1f}%)", level=4, color_hex="#64748b")
                        
                        for i, t in enumerate(filtered_tasks, 1):
                            p = doc.add_paragraph()
                            p.add_run(f"Tarea {i}: ").bold = True
                            p.add_run(f"{t.name}\n")
                            p.add_run("Estado: ").bold = True
                            p.add_run(f"{t.status} ({t.progress:.1f}%)\n")
                            
                            res_names = ", ".join([r.name for r in t.responsibles]) if t.responsibles else "Sin asignar"
                            p.add_run("Responsables: ").bold = True
                            p.add_run(f"{res_names}\n")
                            
                            if t.observations:
                                p.add_run("Observaciones: ").bold = True
                                p.add_run(f"{t.observations}\n")
                            
                            # Evidences
                            evidences = db.query(Evidence).filter(Evidence.task_id == t.id).all()
                            if evidences:
                                p.add_run("Evidencias:\n").bold = True
                                for k, ev in enumerate(evidences, 1):
                                    desc = f" ({ev.description})" if ev.description else ""
                                    p.add_run(f"  {k}. {ev.url}{desc}\n")
                            
                            # Divider
                            doc.add_paragraph("-" * 50)
                            
        # --- EXPORT ---
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio
